"""
Generate Learn2Hire technical documentation as a black-and-white Word document.
Run from repo root: python docs/generate_project_documentation.py
Outputs: docs/Learn2Hire_Project_Documentation.docx and docs/bw_figures/*.png
"""
from __future__ import annotations

import textwrap
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent.parent
OUT_DOC = ROOT / "docs" / "Learn2Hire_Project_Documentation.docx"
FIG_DIR = ROOT / "docs" / "bw_figures"


def make_bw_figure(path: Path, title: str, subtitle: str = "") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 920, 520
    img = Image.new("L", (w, h), 255)
    draw = ImageDraw.Draw(img)
    draw.rectangle([4, 4, w - 5, h - 5], outline=0, width=4)
    draw.line([4, 48, w - 5, 48], fill=0, width=2)
    try:
        font_title = ImageFont.truetype("arial.ttf", 28)
        font_sub = ImageFont.truetype("arial.ttf", 18)
        font_small = ImageFont.truetype("arial.ttf", 14)
    except OSError:
        font_title = font_sub = font_small = ImageFont.load_default()
    draw.text((20, 12), "Learn2Hire — UI placeholder (grayscale)", fill=0, font=font_small)
    y = 70
    for line in textwrap.wrap(title, width=42):
        draw.text((24, y), line, fill=0, font=font_title)
        y += 34
    if subtitle:
        y += 8
        for line in textwrap.wrap(subtitle, width=58):
            draw.text((24, y), line, fill=0, font=font_sub)
            y += 24
    note = (
        "Replace this figure with a full-window screenshot. "
        "Print or export the final report in grayscale for black-and-white distribution."
    )
    ny = h - 72
    for line in textwrap.wrap(note, width=72):
        draw.text((24, ny), line, fill=80, font=font_small)
        ny += 18
    img.save(path, format="PNG")


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_code_block(doc: Document, lines: str):
    p = doc.add_paragraph()
    run = p.add_run(lines.rstrip("\n"))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            table.rows[ri + 1].cells[ci].text = cell
    doc.add_paragraph()


SCREENS = [
    ("01_landing", "Landing page", "Route: / — Marketing home, session cleared on mount."),
    ("02_login", "Login", "Route: /login — JWT stored in localStorage."),
    ("03_signup", "Signup", "Route: /signup — OTP email flow when SMTP configured."),
    ("04_forgot_password", "Forgot password", "Route: /forgot-password — OTP / reset flow."),
    ("05_dashboard_student", "Student dashboard", "Route: /dashboard — Sidebar: Dashboard, Profile, Assessments, Jobs, Learning, Progress."),
    ("06_dashboard_faculty", "Faculty dashboard", "Route: /dashboard — Manage learning, assessments, progress; approval gating on API."),
    ("07_dashboard_company", "Company dashboard", "Route: /dashboard — Job and talent workspace."),
    ("08_dashboard_college", "College dashboard", "Route: /dashboard — Roster, imports, faculty approvals."),
    ("09_dashboard_admin", "Admin dashboard", "Route: /dashboard — Platform analytics, users, colleges."),
    ("10_dashboard_alumni", "Alumni workspace", "Route: /dashboard — Jobs + learning links (lightweight shell)."),
    ("11_learning_public", "Learning hub (public)", "Route: /learning — Catalog and filters without auth."),
    ("12_learning_subject_public", "Learning subject (public)", "Route: /learning/subject/:categorySlug"),
    ("13_material_detail", "Material detail", "Routes: /learning/topic/:slug and /dashboard/learning/topic/:slug"),
    ("14_learning_progress", "My learning progress", "Routes: /learning/progress, /dashboard/learning/progress"),
    ("15_learning_manage", "Learning management (faculty)", "Route: /dashboard/learning/manage — CRUD, image upload, sheet import."),
    ("16_learning_dashboard", "Learning hub (dashboard)", "Route: /dashboard/learning — Authenticated hub shell."),
    ("17_notifications", "Notifications", "Route: /notifications — In-app notification center."),
    ("18_assessments_list", "Assessments list", "Route: /assessments — Published assessments."),
    ("19_assessment_create", "Create assessment", "Route: /assessments/create — Faculty authoring."),
    ("20_assessment_take", "Take assessment", "Route: /assessments/:id — Student attempt / view."),
    ("21_jobs", "Jobs board", "Route: /jobs — Listings, save, suggestions."),
    ("22_job_detail", "Job detail", "Route: /jobs/:id — Detail, apply, JD download when present."),
    ("23_admin_jobs", "Admin jobs", "Route: /admin/jobs — Administrative job oversight."),
    ("24_admin_user", "Admin user profile", "Route: /admin/users/:userId"),
    ("25_admin_college", "Admin college detail", "Route: /admin/colleges/:collegeId"),
    ("26_company_jobs", "Company jobs", "Route: /company/jobs — Employer job management."),
    ("27_company_talent", "Company talent search", "Route: /company/talent — Candidate discovery."),
    ("28_learner_summary", "Learner summary", "Route: /dashboard/learners/:userId — Viewer summary for a user."),
]

# /dashboard views by role (figures 05–10) plus authenticated learning hub (16); gallery at front of doc.
DASHBOARD_GALLERY = list(SCREENS[4:10]) + [SCREENS[15]]


def main():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    for stem, title, sub in SCREENS:
        fig_path = FIG_DIR / f"{stem}.png"
        if not fig_path.exists():
            make_bw_figure(fig_path, title, sub)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Learn2Hire\nTechnical & Functional Specification")
    r.bold = True
    r.font.size = Pt(22)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "Placement and learning platform: student profiles, assessments, jobs, "
        "notifications, and a public or authenticated learning hub."
    )

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Figures include application screenshots (see docs/bw_figures/; refresh via Playwright per README).\n"
    )
    meta.add_run("Generated: April 13, 2026\n")
    meta.add_run("Source repository: Learn2Hire (frontend + backend)")

    doc.add_page_break()

    doc.add_heading("Table of contents", level=1)
    add_para(
        doc,
        "In Microsoft Word: References → Table of Contents → Automatic Table 1. "
        "This file uses built-in Heading styles so the TOC can be generated or refreshed on demand.",
        italic=True,
    )

    doc.add_page_break()
    doc.add_heading("Role dashboards (screenshots)", level=1)
    add_para(
        doc,
        "The following captures show the main /dashboard experience after sign-in for each primary role, "
        "plus the authenticated learning hub (/dashboard/learning). "
        "PNGs live under docs/bw_figures/ (e.g. 05_dashboard_student.png … 10_dashboard_alumni.png, "
        "16_learning_dashboard.png). Regenerate with: cd docs && npm run screenshots "
        "(requires backend, frontend, and npm run seed:doc-users in backend — see repository README).",
    )
    for stem, title, sub in DASHBOARD_GALLERY:
        doc.add_heading(title, level=2)
        add_para(doc, sub, italic=True)
        fig_path = FIG_DIR / f"{stem}.png"
        if fig_path.exists():
            doc.add_picture(str(fig_path), width=Inches(5.9))
        else:
            add_para(doc, f"[Missing image: {fig_path.name} — run documentation screenshot pipeline.]", italic=True)
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("1. Executive summary", level=1)
    add_para(
        doc,
        "Learn2Hire is a web application that connects students, faculty, colleges, companies, "
        "and platform administrators around placement activities and structured learning content. "
        "The system provides REST APIs backed by MongoDB, a React single-page application served "
        "by Vite, and role-aware dashboards that gate features by user role and approval status.",
    )

    doc.add_heading("2. Product scope", level=1)
    add_para(
        doc,
        "In scope: user registration and authentication (including email OTP for signup and "
        "password reset when SMTP is configured), student profile and skills, assessments and "
        "submissions, job postings and applications, saved jobs and suggestions, in-app "
        "notifications, learning categories/subjects/materials with reading progress, faculty "
        "management of study materials (including spreadsheet import and image-assisted creation), "
        "college roster and faculty approval workflows, and administrative analytics and user "
        "management including college account approval.",
    )
    add_para(
        doc,
        "The public learning hub allows anonymous browsing of subjects and materials; personalized "
        "recommendations and progress sync require a signed-in student (or eligible) account.",
    )

    doc.add_heading("3. User roles", level=1)
    add_table(
        doc,
        ["Role", "Primary capabilities"],
        [
            ["student", "Profile, cohort fields, assessments, jobs, learning progress, notifications."],
            ["alumni", "Jobs access; public learning; simplified dashboard shell."],
            ["faculty", "Assessments, submissions overview, learning content management (when approved)."],
            ["company", "Post jobs, manage applications, talent search, express interest."],
            ["college", "Roster, import students, pending faculty list, approval actions (when approved)."],
            ["admin", "Platform analytics, user directory, role changes, college lifecycle, imports."],
        ],
    )

    doc.add_heading("4. High-level architecture", level=1)
    add_para(
        doc,
        "Three logical tiers: (1) Browser client — React Router, localStorage session, fetch to /api; "
        "(2) API tier — Express on Node.js, JWT middleware, controllers; (3) Data tier — MongoDB via Mongoose. "
        "Static uploads (e.g. profile photos, learning images) are served from the backend /uploads path.",
    )
    add_code_block(
        doc,
        """
+------------------+       HTTPS/HTTP        +-------------------+
|   Web browser    | <----------------------> |  Vite dev :3000   |
|   (React SPA)    |      same-origin /api    |  (proxies /api)   |
+------------------+ +---------+---------+
                                                        |
                                                        v
                                              +-------------------+
                                              |  Express API :5000|
                                              |  /api/* JSON |
                                              +---------+---------+
                                                        |
                                                        v
                                              +-------------------+
                                              | MongoDB       |
                                              +-------------------+
""",
    )

    doc.add_heading("5. Technology stack", level=1)
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ["Frontend", "React 18, React Router 6, Vite 6, Tailwind CSS 4"],
            ["Backend", "Node.js, Express 5, Mongoose, JWT, bcrypt"],
            ["Data", "MongoDB"],
            ["Email / OTP", "Nodemailer (when SMTP env configured); dev OTP echo optional"],
        ],
    )

    doc.add_heading("6. Runtime and configuration", level=1)
    add_para(
        doc,
        "Local development: backend listens on port 5000 by default; frontend Vite server on 3000. "
        "The Vite config proxies requests beginning with /api to the backend (override with VITE_PROXY_TARGET). "
        "Backend requires .env with MongoDB connection and JWT secret; optional SMTP variables for OTP mail.",
    )

    doc.add_heading("7. Security model", level=1)
    add_para(
        doc,
        "Authenticated requests send Authorization: Bearer <token>. The protect middleware verifies JWT, "
        "loads the User document, and enforces:",
    )
    doc.add_paragraph("Faculty: if facultyApprovalStatus is pending or rejected, most API paths return 403; "
                      "/api/auth/me and profile photo routes remain available where coded.", style="List Bullet")
    doc.add_paragraph("College: if collegeApprovalStatus is pending or rejected, only /api/auth/me is allowed.", style="List Bullet")
    doc.add_paragraph("Admin: non-builtin admin emails are restricted to /api/auth/me only.", style="List Bullet")
    add_para(
        doc,
        "Optional JWT (optionalProtect) is used on some learning GET routes so the same endpoint can return "
        "anonymous or personalized payloads.",
    )

    doc.add_heading("8. Data entities (MongoDB collections)", level=1)
    add_table(
        doc,
        ["Model", "Purpose"],
        [
            ["User", "Identity, role, passwords (select:false), approval flags, college affiliation, profile photo path."],
            ["StudentProfile", "Extended student fields, skills, cohort, visibility to companies."],
            ["Assessment", "Assessment definitions for faculty/students."],
            ["AssessmentSubmission", "Student attempts and scores."],
            ["Job", "Job postings; attachments metadata for JD."],
            ["JobApplication", "Applications and pipeline status."],
            ["SavedJob", "Student saved jobs."],
            ["JobStudentInterest", "Interest signals between students and companies."],
            ["Notification", "In-app notifications."],
            ["StudyMaterial", "Learning content units (slug, subject, body, media)."],
            ["Subject", "Subject master data."],
            ["LearningCategory", "Legacy/public categories for catalog."],
            ["LearningProgress", "Per-user progress for materials."],
            ["EmailOtp", "OTP records for signup and password reset."],
        ],
    )

    doc.add_heading("9. Data flow diagrams", level=1)
    doc.add_heading("9.1 Context (Level 0)", level=2)
    add_code_block(
        doc,
        """
                    +-----------+
                    |  Actors   |
                    | (student, |
                    | faculty,  |
                    | company,  |
                    | college,  |
                    | admin,    |
                    | visitor)  |
                    +-----+-----+
                          |
          browse / JSON |
 v
                   +--------------+
                   |  Learn2Hire  |
                   | System     |
                   +------+-------+
                          |
              persist |
                          v
                   +--------------+
                   |   MongoDB    |
                   +--------------+
""",
    )

    doc.add_heading("9.2 Authentication flow", level=2)
    add_code_block(
        doc,
        """
 [Client] --POST /api/auth/login--> [AuthController] --verify bcrypt--> [User]
 |                                      |
     |<------- JWT + user JSON -------------+
     |
 localStorage: token, user
     |
     v Subsequent calls: Header Authorization: Bearer <token>
     |
     v
 [protect middleware] --jwt.verify--> attach req.user --> route handler
""",
    )

    doc.add_heading("9.3 Assessment lifecycle", level=2)
    add_code_block(
        doc,
        """
 [Faculty] --> POST /api/assessments --> [Assessment document]
 [Student] --> GET /api/assessments --> list published
 [Student] --> GET /api/assessments/:id --> load quiz
 [Student] --> POST /api/submissions --> [AssessmentSubmission]
 [Faculty] --> GET /api/submissions/assessment/:id --> review cohort results
""",
    )

    doc.add_heading("9.4 Job application flow", level=2)
    add_code_block(
        doc,
        """
 [Company] --> POST /api/jobs --> Job
 [Student] --> GET /api/jobs / suggestions / saved
 [Student] --> POST /api/jobs/:id/apply --> JobApplication
 [Company] --> GET /api/jobs/:id/applications --> pipeline
 PATCH /api/jobs/applications/:applicationId/status --> status updates
""",
    )

    doc.add_heading("9.5 Learning progress flow", level=2)
    add_code_block(
        doc,
        """
 [Any] --> GET /api/learning/materials (optional JWT)
 [Student JWT] --> GET /api/learning/progress/me
 [Student JWT] --> PUT /api/learning/progress/material/:slug --> LearningProgress
 [Faculty JWT] --> /api/learning/manage/* --> CRUD StudyMaterial / categories / import
""",
    )

    doc.add_heading("9.6 College / faculty governance", level=2)
    add_code_block(
        doc,
        """
 [College JWT] --> GET /api/college/faculty/pending
 [College JWT] --> PATCH /api/college/faculty/:id/approval
 [Admin JWT] --> PATCH /api/admin/colleges/:id/approval
 Faculty API access requires facultyApprovalStatus approved (middleware).
""",
    )

    doc.add_heading("10. REST API mount points", level=1)
    add_table(
        doc,
        ["Mount path", "Router module", "Notes"],
        [
            ["/api/auth", "authRoutes.js", "OTP, signup, login, me"],
            ["/api/profile", "profileRoutes.js", "Protected profile CRUD, photo, skills, summary by userId"],
            ["/api/assessments", "assessmentRoutes.js", "Protected CRUD"],
            ["/api/submissions", "submissionRoutes.js", "Protected submit and list"],
            ["/api/jobs", "jobRoutes.js", "Protected jobs, apply, save, company dashboard, talent"],
            ["/api/admin", "adminRoutes.js", "Protected analytics, users, colleges, imports"],
            ["/api/notifications", "notificationRoutes.js", "Protected"],
            ["/api/learning", "learningRoutes.js", "Public reads + optional auth; /manage under protect"],
            ["/api/subjects", "subjectRoutes.js", "Protected subject CRUD"],
            ["/api/college", "collegeRoutes.js", "Protected college operations"],
            ["/api/health", "server.js", "Health JSON"],
        ],
    )

    doc.add_heading("11. Frontend routing map", level=1)
    add_para(doc, "Public routes (no JWT required for initial page load):", bold=True)
    pub_rows = [
        ["/", "LandingPage"],
        ["/login", "Login"],
        ["/signup", "Signup"],
        ["/forgot-password", "ForgotPassword"],
        ["/learning", "LearningHomePage public mode"],
        ["/learning/topic/:slug", "MaterialDetailsPage"],
        ["/learning/subject/:categorySlug", "LearningSubjectPage public mode"],
        ["/learning/:categorySlug", "Redirect to /learning/subject/:slug"],
    ]
    add_table(doc, ["Path", "Component / behavior"], pub_rows)

    add_para(doc, "Protected routes (ProtectedRoute redirects to /login):", bold=True)
    prot_rows = [
        ["/dashboard", "Dashboard role router"],
        ["/dashboard/learners/:userId", "LearnerSummaryPage"],
        ["/learning/progress", "MyLearningProgressPage"],
        ["/dashboard/learning/*", "Learning hub, manage, subject, topic, progress"],
        ["/notifications", "NotificationsPage"],
        ["/assessments", "AssessmentsList"],
        ["/assessments/create", "CreateAssessment"],
        ["/assessments/:id", "StudentAssessment"],
        ["/jobs", "JobsPage"],
        ["/jobs/:id", "JobDetailsPage"],
        ["/admin/jobs", "AdminJobsPage"],
        ["/admin/users/:userId", "AdminUserProfilePage"],
        ["/admin/colleges/:collegeId", "AdminCollegeDetailPage"],
        ["/company/jobs", "CompanyJobsPage"],
        ["/company/talent", "CompanyTalentPage"],
    ]
    add_table(doc, ["Path", "Component"], prot_rows)
    add_para(doc, "Unknown paths navigate to / (catch-all).", italic=True)

    doc.add_heading("12. Dashboard composition by role", level=1)
    add_para(
        doc,
        "The /dashboard route reads role from localStorage user object and renders: StudentDashboard, "
        "FacultyDashboard, CompanyDashboard, CollegeDashboard, AdminDashboard, a dedicated alumni shell, "
        "or a generic 'workspace setup' placeholder for unexpected roles.",
    )
    add_para(doc, "Student sidebar sections (in-app): Dashboard, Profile, Assessments, external link to Jobs, Learning, Progress.", bold=True)
    add_para(
        doc,
        "Student dashboard data pulls include profile, assessments, submissions, suggested jobs, "
        "recommended materials, and aggregated learning summary metrics (started, completed, in progress, time, average progress).",
    )
    add_para(doc, "Faculty sidebar: Dashboard, Manage learning (route to /dashboard/learning/manage), Profile, Assessments, Progress.", bold=True)
    add_para(
        doc,
        "Admin dashboard includes metric cards, directory sections for colleges/companies/faculty, collapsible "
        "panels for imports and user operations, and deep links to admin user and college detail routes.",
    )
    add_para(
        doc,
        "Company dashboard loads /api/jobs/company/dashboard and surfaces metrics (total jobs, open jobs, "
        "applications, shortlisted counts), recent jobs, recent applications, inline job creation with draft/open "
        "workflow, and application status updates via PATCH /api/jobs/applications/:applicationId/status.",
        bold=True,
    )
    add_para(
        doc,
        "College dashboard aggregates /api/college/insights, roster listing, student import, pending faculty "
        "approvals, and contextual links to assessments and jobs relevant to the campus view.",
        bold=True,
    )
    add_para(
        doc,
        "Notifications: GET /api/notifications and /api/notifications/unread-count; mark one read with "
        "PATCH /api/notifications/:id/read; mark all with PATCH /api/notifications/read-all. The React tree "
        "wraps NotificationProvider for client-side coordination with the notifications page.",
    )

    doc.add_heading("13. Screen catalog and figure index", level=1)
    add_para(
        doc,
        "Figures under docs/bw_figures/ are normally produced with Playwright (see README: "
        "\"Technical documentation\"). Missing files are filled with grayscale placeholders when this script runs. "
        "For strictly black-and-white printing, use Word’s Picture → Color → Grayscale.",
    )
    for stem, title, sub in SCREENS:
        doc.add_heading(f"Figure — {title}", level=2)
        add_para(doc, sub, italic=True)
        pic_path = FIG_DIR / f"{stem}.png"
        doc.add_picture(str(pic_path), width=Inches(5.8))
        doc.add_paragraph()

    doc.add_heading("14. Operational procedures", level=1)
    doc.add_paragraph("Seed learning content (from backend/): npm run seed:subjects; npm run seed:materials; npm run seed:learning; verify with npm run verify:learning.", style="List Number")
    doc.add_paragraph("Builtin admin emails are ensured on server start (ensureBuiltinAdmins).", style="List Number")
    doc.add_paragraph("API responses under /api use Cache-Control: no-store to avoid stale cached JSON in browsers.", style="List Number")

    doc.add_heading("15. Appendix — workflow checklist (printable)", level=1)
    add_para(doc, "Student onboarding: Signup → verify OTP → login → complete profile and cohort → browse learning → take assessments → apply to jobs.", bold=True)
    add_para(doc, "Faculty onboarding: Signup with college affiliation → pending approval → after approval, create assessments and manage learning materials.", bold=True)
    add_para(doc, "Employer: Company login → create jobs → review applications → update statuses → use talent search as implemented.", bold=True)
    add_para(doc, "College: Approve faculty → maintain roster → import students via sheet endpoints as exposed in UI.", bold=True)
    add_para(doc, "Admin: Approve colleges → manage users and roles → monitor analytics → curate platform data.", bold=True)

    doc.save(OUT_DOC)
    print(f"Wrote {OUT_DOC}")


if __name__ == "__main__":
    main()
